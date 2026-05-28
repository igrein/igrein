import io
import logging
import textwrap
from datetime import datetime
from typing import List, Dict
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.utils import ImageReader
from docx import Document as DocxDocument
from docx.shared import Pt
import os
import re
import tempfile
from docx.shared import Inches
import logging
from docx.shared import Pt, Inches
from docx.enum.table import WD_ALIGN_VERTICAL
from document_models import VariantDocument, TextBlock, FormulaBlock, TableBlock

logger = logging.getLogger(__name__)

try:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    MATPLOTLIB_AVAILABLE = True
except ImportError:
    MATPLOTLIB_AVAILABLE = False
    logging.warning("Matplotlib не установлен, LaTeX в PDF не будет рендериться")

logger = logging.getLogger(__name__)


# Загрузка шрифта
def find_font():
    """Ищет шрифт с поддержкой кириллицы в системе"""
    possible_paths = [
        # Путь локальный
        "/home/igrein/hakaton/taskforge_backend/fonts/DejaVuSans.ttf",
        "fonts/DejaVuSans.ttf",
        "DejaVuSans.ttf",
        # Windows
        "C:\\Windows\\Fonts\\Arial.ttf",
        "C:\\Windows\\Fonts\\Times.ttf",
        # Linux
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        # macOS
        "/System/Library/Fonts/Arial.ttf",
    ]
    
    for path in possible_paths:
        if os.path.exists(path):
            logger.info(f"✅ Шрифт найден: {path}")
            return path
    logger.warning("❌ Шрифт с кириллицей не найден")
    return None


# Регистрируем шрифт
FONT_PATH = find_font()
FONT_AVAILABLE = FONT_PATH is not None

if FONT_AVAILABLE:
    try:
        pdfmetrics.registerFont(TTFont("UnicodeFont", FONT_PATH))
        logger.info("✅ Шрифт UnicodeFont зарегистрирован")
    except Exception as e:
        logger.warning(f"Ошибка регистрации шрифта: {e}")
        FONT_AVAILABLE = False


def clean_latex_mess(text: str) -> str:
    """Очищает текст от мусора, который генерирует LLM"""
    # Убираем экранированные пробелы
    text = text.replace(r'\,', ' ')
    text = text.replace(r'\ ', ' ')
    
    # Убираем обратные слеши перед обычными словами
    text = re.sub(r'\\([a-zA-Z]+)', r'\1', text)
    
    # Убираем лишние пробелы
    text = re.sub(r'\s+', ' ', text)
    
    return text.strip()


def normalize_text_spacing(text: str) -> str:
    """Нормализует переносы строк для компактного отображения"""
    import re
    
    # Убираем тройные+ переносы
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # Убираем переносы между "А)" и словом
    text = re.sub(r'([А-ЯA-Z]\))\s*\n+\s*', r'\1 ', text)
    
    # Убираем переносы между "1." и словом
    text = re.sub(r'(\d+\.)\s*\n+\s*', r'\1 ', text)
    
    # Убираем пустые строки вокруг маркеров
    text = re.sub(r'\n+([А-Я]\))\s*\n+', r'\n\1 ', text)
    text = re.sub(r'\n+(\d+\.)\s*\n+', r'\n\1 ', text)
    
    return text.strip()


def export_to_pdf(variants: List[VariantDocument], answers: Dict[int, str] = None, include_answers: bool = False) -> io.BytesIO:
    """Экспорт в PDF с возможностью включить ответы учителя"""
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    
    buffer = io.BytesIO()
    
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
    styles = getSampleStyleSheet()
    
    # Стиль для кириллицы
    if FONT_AVAILABLE:
        styles.add(ParagraphStyle(name='RussianNormal', parent=styles['Normal'], fontName='UnicodeFont', fontSize=11, leading=14))
        styles.add(ParagraphStyle(name='RussianHeading', parent=styles['Heading1'], fontName='UnicodeFont', fontSize=14, leading=18))
        styles.add(ParagraphStyle(name='RussianTitle', parent=styles['Title'], fontName='UnicodeFont', fontSize=16, leading=20))
        styles.add(ParagraphStyle(name='RussianAnswer', parent=styles['Normal'], fontName='UnicodeFont', fontSize=10, leading=12, textColor=colors.green))
        
        normal_style = styles['RussianNormal']
        heading_style = styles['RussianHeading']
        title_style = styles['RussianTitle']
        answer_style = styles['RussianAnswer']
    else:
        normal_style = styles['Normal']
        heading_style = styles['Heading1']
        title_style = styles['Title']
        answer_style = styles['Normal']
    
    story = []
    
    story.append(Paragraph("TaskForge AI - Варианты заданий", title_style))
    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph(f"Сгенерировано: {datetime.now().strftime('%d.%m.%Y %H:%M')}", normal_style))
    if include_answers:
        story.append(Paragraph("📌 Включены ответы для учителя", normal_style))
    story.append(Spacer(1, 0.3 * inch))
    
    for variant in variants:
        story.append(Paragraph(f"ВАРИАНТ {variant.number}", heading_style))
        story.append(Spacer(1, 0.1 * inch))
        
        for block in variant.blocks:
            if block.type == "text":
                clean_text = normalize_text_spacing(block.content)
                if clean_text:
                    html_text = clean_text.replace('\n', '<br/>')
                    story.append(Paragraph(html_text, normal_style))
                    story.append(Spacer(1, 0.05 * inch))
            
            elif block.type == "formula":
                if block.latex:
                    story.append(Paragraph(f"<i>{block.latex}</i>", normal_style))
                    story.append(Spacer(1, 0.05 * inch))
            
            elif block.type == "table":
                headers, rows = parse_latex_table_robust(block.latex)
                if headers and rows:
                    table_data = [headers] + rows
                    table = Table(table_data)
                    table.setStyle(TableStyle([
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ]))
                    story.append(table)
                    story.append(Spacer(1, 0.1 * inch))
        
        # Добавляем ответ учителя если нужно
        if include_answers and answers and str(variant.number) in answers:
            answer_text = answers[str(variant.number)]
            if answer_text and answer_text != "—":
                story.append(Spacer(1, 0.05 * inch))
                story.append(Paragraph(f"🔑 <b>Ответ для учителя:</b> {answer_text}", answer_style))
                story.append(Spacer(1, 0.1 * inch))
        
        story.append(Spacer(1, 0.15 * inch))
    
    doc.build(story)
    buffer.seek(0)
    return buffer


def export_to_docx(variants: List[VariantDocument], answers: Dict[int, str] = None, include_answers: bool = False) -> io.BytesIO:
    """Экспорт в DOCX с возможностью включить ответы учителя"""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_LINE_SPACING
    
    buffer = io.BytesIO()
    doc = DocxDocument()
    
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    
    title = doc.add_heading("TaskForge AI", level=0)
    title.alignment = 1
    title.paragraph_format.space_after = Pt(6)
    
    doc.add_paragraph(f"Сгенерировано: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    if include_answers:
        doc.add_paragraph("📌 Включены ответы для учителя")
    
    for variant in variants:
        heading = doc.add_heading(f"ВАРИАНТ {variant.number}", level=1)
        heading.paragraph_format.space_before = Pt(10)
        heading.paragraph_format.space_after = Pt(4)
        
        for block in variant.blocks:
            if block.type == "text":
                clean_text = normalize_text_spacing(block.content)
                if not clean_text:
                    continue
                
                para = doc.add_paragraph()
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.line_spacing = 1.0
                
                lines = clean_text.split('\n')
                for i, line in enumerate(lines):
                    if line.strip() or (i > 0 and i < len(lines)-1):
                        run = para.add_run(line)
                        run.font.size = Pt(11)
                        if i < len(lines) - 1:
                            para.add_run().add_break()
            
            elif block.type == "formula":
                if block.latex:
                    para = doc.add_paragraph()
                    run = para.add_run(block.latex)
                    run.italic = True
                    run.font.size = Pt(11)
                    para.paragraph_format.space_before = Pt(2)
                    para.paragraph_format.space_after = Pt(2)
                    para.paragraph_format.line_spacing = 1.0
            
            elif block.type == "table":
                if block.latex:
                    headers, rows = parse_latex_table_robust(block.latex)
                    if headers and rows:
                        table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
                        table.style = 'Light Grid Accent 1'
                        
                        for row in table.rows:
                            for cell in row.cells:
                                cell.paragraphs[0].paragraph_format.space_before = Pt(0)
                                cell.paragraphs[0].paragraph_format.space_after = Pt(0)
                                cell.paragraphs[0].paragraph_format.line_spacing = 1.0
                        
                        for i, header in enumerate(headers):
                            table.cell(0, i).text = str(header).strip()
                        
                        for row_idx, row in enumerate(rows):
                            for col_idx, cell_value in enumerate(row):
                                if col_idx < len(headers):
                                    table.cell(row_idx + 1, col_idx).text = str(cell_value).strip()
        
        # Добавляем ответ учителя если нужно
        if include_answers and answers and str(variant.number) in answers:
            answer_text = answers[str(variant.number)]
            if answer_text and answer_text != "—":
                para = doc.add_paragraph()
                para.paragraph_format.space_before = Pt(4)
                para.paragraph_format.space_after = Pt(2)
                run = para.add_run(f"🔑 Ответ для учителя: {answer_text}")
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 128, 0)
        
        if variant != variants[-1]:
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after = Pt(0)
    
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def parse_latex_table_robust(latex: str) -> tuple:
    """Робостный парсинг LaTeX таблицы"""
    try:
        latex = latex.replace('\n', ' ').replace('\r', '')
        
        import re
        
        # Паттерн для array или tabular
        pattern = r'\\begin\{array\}\{.*?\}(.*?)\\end\{array\}'
        match = re.search(pattern, latex, re.DOTALL)
        
        if not match:
            pattern = r'\\begin\{tabular\}\{.*?\}(.*?)\\end\{tabular\}'
            match = re.search(pattern, latex, re.DOTALL)
        
        if not match:
            return [], []
        
        content = match.group(1)
        
        # Разбиваем на строки
        rows = []
        for row in content.split('\\\\'):
            row = row.strip()
            # Убираем \hline
            row = row.replace('\\hline', '').strip()
            if not row:
                continue
            
            # Разбиваем по & 
            cells = [cell.strip() for cell in row.split('&')]
            # Очищаем каждую ячейку от LaTeX команд
            clean_cells = []
            for cell in cells:
                cell = re.sub(r'\\text\{(.*?)\}', r'\1', cell)
                cell = re.sub(r'\\textbf\{(.*?)\}', r'\1', cell)
                cell = cell.strip()
                clean_cells.append(cell)
            
            if clean_cells:
                rows.append(clean_cells)
        
        if not rows:
            return [], []
        
        # Первая строка - заголовки
        headers = rows[0]
        data_rows = rows[1:] if len(rows) > 1 else []
        
        return headers, data_rows
        
    except Exception as e:
        logger.error(f"Ошибка парсинга таблицы: {e}")
        return [], []

def strip_latex_delimiters(text: str) -> str:
    """Заменяет LaTeX формулы на плейсхолдеры для экспорта"""
    return text



def extract_latex_parts(text: str):
    """Разбивает текст на части:
    - обычный текст
    - LaTeX формулы
    """

    pattern = r'\$\$(.*?)\$\$|\\\((.*?)\\\)|\\\[(.*?)\\\]'

    parts = []

    last_end = 0

    for match in re.finditer(pattern, text, re.DOTALL):

        start, end = match.span()

        # Обычный текст
        if start > last_end:
            parts.append((
                "text",
                text[last_end:start]
            ))

        latex_content = (
            match.group(1)
            or match.group(2)
            or match.group(3)
        )

        parts.append((
            "latex",
            latex_content.strip()
        ))

        last_end = end

    # Остаток текста
    if last_end < len(text):
        parts.append((
            "text",
            text[last_end:]
        ))

    return parts
